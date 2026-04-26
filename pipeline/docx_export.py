from __future__ import annotations

from pathlib import Path
import json
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _add_kv(paragraph, label: str, value: str) -> None:
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value)


def _confidence_reason(step: dict[str, Any]) -> str:
    components = step.get("confidence_components") or {}
    reasons = []
    if step.get("confidence") != "high":
        reasons.append("model or local evidence was not high confidence")
    if step.get("risky"):
        reasons.append("risk filter flagged the wording for review")
    if float(components.get("text_delta", 0) or 0) > 0.6:
        reasons.append("OCR changed substantially around this segment")
    if float(components.get("visual", 0) or 0) < 0.04:
        reasons.append("visual boundary evidence was weak")
    return "; ".join(reasons) or "evidence appears consistent"


def export_docx(
    process_name: str,
    phases: dict[str, list[dict[str, Any]]] | list[dict[str, Any]],
    output_path: str | Path,
    summary: str | None = None,
    target_audience: str = "New employee",
    department_notes: str = "",
    warnings: list[str] | None = None,
    job_metadata: dict[str, Any] | None = None,
    cleanup_report: dict[str, Any] | None = None,
) -> str:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = document.add_heading(f"SOP: {process_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("1. Summary", level=1)
    if isinstance(phases, list):
        phase_sections = [(section.get("phase", "Review and validate"), section.get("steps", [])) for section in phases]
    else:
        phase_sections = list(phases.items())

    step_count = sum(len(items) for _, items in phase_sections)
    quality = (cleanup_report or {}).get("quality_report") or {}
    readiness_label = {
        "demo_ready": "Demo-ready",
        "needs_review": "Needs manual review",
        "not_ready": "Not ready",
    }.get(quality.get("readiness"), "Needs manual review")
    if summary:
        document.add_paragraph(summary)
    else:
        coverage_ratio = quality.get("coverage_ratio_after_cleanup")
        event_segments = quality.get("event_segments")
        coverage_line = (
            f"Coverage: {step_count} cleaned steps from {event_segments} event segments"
            + (f" ({coverage_ratio:.3f})." if isinstance(coverage_ratio, (float, int)) else ".")
            if event_segments
            else f"Coverage: {step_count} evidence-backed steps."
        )
        document.add_paragraph(
            f"This SOP documents the process captured in the uploaded screen recording. "
            f"Readiness: {readiness_label}. {coverage_line}"
        )
    p = document.add_paragraph()
    _add_kv(p, "Target audience", target_audience)
    if department_notes:
        p = document.add_paragraph()
        _add_kv(p, "Department / system notes", department_notes)

    document.add_heading("Prerequisites", level=2)
    document.add_paragraph("Access to the systems shown in the screenshots.")
    document.add_paragraph("Permission to view, export, edit, or post records as required by the business process.")
    document.add_paragraph("Review any low-confidence steps before using this SOP for training or production work.")

    document.add_heading("Assumptions", level=2)
    document.add_paragraph("The SOP is based only on visible screen evidence and OCR extracted from the uploaded recording.")
    document.add_paragraph("Exact values are generalized unless they are clearly field labels or business concepts.")
    if warnings:
        document.add_heading("Evidence warnings", level=2)
        for warning in warnings:
            document.add_paragraph(warning, style=None)

    document.add_heading("2. Steps grouped by phases", level=1)
    for phase_name, steps in phase_sections:
        document.add_heading(phase_name, level=2)
        for step in steps:
            document.add_heading(f"Step {step['step_number']}", level=3)
            p = document.add_paragraph()
            _add_kv(p, "System", step.get("system", "Other"))
            p = document.add_paragraph()
            _add_kv(p, "Action", step.get("action", "Review the visible process screen."))
            p = document.add_paragraph()
            _add_kv(p, "Expected output", step.get("expected_output", "The relevant process information is visible."))
            p = document.add_paragraph()
            _add_kv(p, "Confidence", step.get("confidence", "medium").title())
            if step.get("risky"):
                p = document.add_paragraph()
                _add_kv(p, "Review note", "This step was flagged for conservative review.")

            screenshot = step.get("screenshot")
            if screenshot and Path(screenshot).exists():
                start = float(step.get("start_time_sec", step.get("time_sec", 0)) or 0)
                end = float(step.get("end_time_sec", step.get("time_sec", 0)) or 0)
                state = step.get("screen_state_id", "unknown")
                document.add_paragraph(f"Screenshot evidence: {start:.1f}s-{end:.1f}s, screen state {state}")
                try:
                    document.add_picture(str(screenshot), width=Inches(6.2))
                except Exception:
                    document.add_paragraph(f"Screenshot could not be embedded: {screenshot}")

    document.add_heading("3. Screenshots per step", level=1)
    document.add_paragraph("Screenshots are embedded under each step as supporting evidence.")

    document.add_heading("4. Confidence indicators", level=1)
    document.add_paragraph("High: clear visual and text evidence supports the step.")
    document.add_paragraph("Medium: the action is likely based on visible context but may be generalized.")
    document.add_paragraph("Low: evidence is limited; review the screenshot before using the step operationally.")

    low_confidence = [step for _, steps in phase_sections for step in steps if step.get("confidence") != "high"]
    document.add_heading("Low-confidence review checklist", level=1)
    if low_confidence:
        for step in low_confidence:
            document.add_paragraph(
                f"Step {step.get('step_number')}: review the screenshot and confirm the action/output wording "
                f"({_confidence_reason(step)})."
            )
    else:
        document.add_paragraph("No low-confidence steps were identified.")

    document.add_heading("Appendix: Job metadata", level=1)
    if job_metadata:
        for key, value in job_metadata.items():
            if key in {"quality_profile", "cost_estimate"}:
                document.add_paragraph(f"{key}: {json.dumps(value, ensure_ascii=False)}")
            elif key not in {"warnings"}:
                document.add_paragraph(f"{key}: {value}")
    else:
        document.add_paragraph("No job metadata was recorded.")

    if cleanup_report:
        document.add_heading("Cleanup and quality report", level=1)
        document.add_paragraph(f"Original step count: {quality.get('step_count_before', 'unknown')}")
        document.add_paragraph(f"Cleaned step count: {quality.get('step_count_after', 'unknown')}")
        document.add_paragraph(f"Event segments: {quality.get('event_segments', 'unknown')}")
        document.add_paragraph(
            f"Pre-cleanup coverage ratio: {quality.get('coverage_ratio_before_cleanup', 'unknown')}"
        )
        document.add_paragraph(
            f"Final coverage ratio: {quality.get('coverage_ratio_after_cleanup', 'unknown')}"
        )
        document.add_paragraph(f"Removed steps: {quality.get('removed_count', 0)}")
        document.add_paragraph(f"Merged steps: {quality.get('merged_count', 0)}")
        document.add_paragraph(f"Passive filler removed: {quality.get('passive_filler_removed_count', 0)}")
        document.add_paragraph(f"Chronological order valid: {quality.get('chronological_order_valid', 'unknown')}")
        document.add_paragraph(f"Chronological violations: {quality.get('chronological_violations_count', 'unknown')}")
        document.add_paragraph(f"Quality score: {quality.get('quality_score', 'unknown')}")
        document.add_paragraph(f"Readiness: {quality.get('readiness', 'unknown')}")
        for warning in quality.get("warnings", []) or []:
            document.add_paragraph(f"Warning: {warning}")
        for blocker in quality.get("readiness_blockers", []) or []:
            document.add_paragraph(f"Readiness blocker: {blocker}")

        removed = cleanup_report.get("removed_steps") or []
        if removed:
            document.add_heading("Removed steps", level=2)
            for item in removed:
                document.add_paragraph(
                    f"Original step {item.get('original_step_number')}: {item.get('action')} "
                    f"Reason: {item.get('reason')}"
                )

        merged = cleanup_report.get("merged_steps") or []
        if merged:
            document.add_heading("Merged steps", level=2)
            for item in merged:
                document.add_paragraph(
                    f"Source steps {item.get('source_step_numbers')}: {item.get('merged_action')} "
                    f"Reason: {item.get('reason')}"
                )

    document.save(output_path)
    return str(output_path)
